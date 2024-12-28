import cv2
import os
import numpy as np
import re
import fitz
from PIL import Image
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.platypus import Image as Img
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.pdfgen import canvas
import time
import shutil
import threading
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json










STATE_FILE = "state.json"

def save_state(pdf_path, persist_state):
    """
    Saves the current state (pdf path and persist state) to a JSON file.

    This function writes the provided pdf path and persist state to a file named `STATE_FILE`. 
    It is used for saving the current context to persist data across sessions.

    Args:
        pdf_path (str): The path to the PDF file to be saved.
        persist_state (dict): A dictionary containing the state to be persisted (e.g., user preferences or settings).
    
    Returns:
        None
    """
    state = {
        "pdf_path": pdf_path,
        "persist_state": persist_state
    }
    with open(STATE_FILE, "w") as f:
        json.dump(state, f)  

def load_state():
    """
    Loads the state from the saved JSON file.

    This function attempts to read the state file (`STATE_FILE`) and retrieve the saved 
    pdf path and persist state. If the file doesn't exist or there is an error, 
    an exception is raised.

    Returns:
        tuple: A tuple containing the pdf path (str) and persist state (dict).
    
    Raises:
        ValueError: If the state file is not found or is corrupted (JSON decode error).
    """
    try:
        with open(STATE_FILE, "r") as f:
            state = json.load(f)
        pdf_path = state.get("pdf_path")
        persist_state = state.get("persist_state")
        return pdf_path, persist_state
    except FileNotFoundError:
        raise ValueError("State file not found.")
    except json.JSONDecodeError:
        raise ValueError("Error decoding state file.")

def clear_state():
    """
    Clears the saved state by removing the state file.

    This function checks if the state file exists and deletes it. 
    It is used to reset or clear the saved state.

    Returns:
        None
    """
    if os.path.exists(STATE_FILE):
        os.remove(STATE_FILE)
        print("\nCurrent State Cleared\n")

def detectColor(img, hsv, show_process=False):
    """
    Detects a specific color range in the provided image and returns the result.

    This function converts the input image to HSV color space and creates a mask for the 
    specified color range. It then applies the mask to the image to extract the region 
    that matches the color. Optionally, intermediate images can be displayed during the process.

    Args:
        img (np.ndarray): The input image in BGR color space (NumPy array).
        hsv (list): A list of six integers representing the HSV color range: 
                    [H_min, H_max, S_min, S_max, V_min, V_max].
        show_process (bool, optional): If True, intermediate images (HSV and result) will be shown. Default is False.

    Returns:
        tuple: A tuple containing the result image (imgResult) and the HSV image (imgHSV).
    """
    imgHSV = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
    cv2.imshow("HSV", imgHSV) if show_process else ""
    lower = np.array([hsv[0], hsv[2], hsv[4]])
    upper = np.array([hsv[1], hsv[3], hsv[5]])
    mask = cv2.inRange(imgHSV, lower, upper)
    imgResult = cv2.bitwise_and(img, img, mask=mask)
    cv2.imshow("Result", imgResult) if show_process else ""
    return imgResult, imgHSV


def getContours(img, imgDraw, showCanny=False, minArea=1000, filter=0, cThr=[100, 100], draw=True):
    """
    Detects contours in an image and optionally draws bounding boxes around them.

    This function processes the input image through various steps (grayscale conversion, Gaussian blur, 
    Canny edge detection, dilation, and morphological closing) to find contours. The contours are filtered 
    based on area and number of polygon vertices, and bounding boxes are drawn if specified.

    Args:
        img (np.ndarray): The input image in BGR color space (NumPy array).
        imgDraw (np.ndarray): The image on which contours will be drawn.
        showCanny (bool, optional): If True, the Canny edge detection result will be shown. Default is False.
        minArea (int, optional): The minimum area of the contours to consider. Default is 1000.
        filter (int, optional): The number of vertices of the polygonal approximation to filter. Default is 0 (no filtering).
        cThr (list, optional): The thresholds for Canny edge detection in the form [low_threshold, high_threshold]. Default is [100, 100].
        draw (bool, optional): If True, bounding boxes will be drawn around the detected contours. Default is True.

    Returns:
        tuple: A tuple containing:
            - imgDraw (np.ndarray): The image with bounding boxes drawn around the contours.
            - finalContours (list): A list of contours, each represented as a list containing:
              [number of vertices, area, contour approximation, bounding box, contour].
    """
    imgDraw = imgDraw.copy()
    imgGray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    imgBlur = cv2.GaussianBlur(imgGray, (5, 5), 1)
    imgCanny = cv2.Canny(imgBlur, cThr[0], cThr[1])
    kernel = np.ones((10, 10), np.uint8)
    imgDial = cv2.dilate(imgCanny, kernel, iterations=1)
    imgClose = cv2.morphologyEx(imgDial, cv2.MORPH_CLOSE, kernel)

    if showCanny: cv2.imshow("Canny", imgClose)

    contours, hierarchy = cv2.findContours(imgClose, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    finalContours = []
    for i in contours:
        area = cv2.contourArea(i)
        if area > minArea:
            peri = cv2.arcLength(i, True)
            approx = cv2.approxPolyDP(i, 0.02*peri, True)
            bbox = cv2.boundingRect(approx)
            if filter > 0: 
                if len(approx) == filter:
                    finalContours.append([len(approx), area, approx, bbox, i])
            else:
                finalContours.append([len(approx), area, approx, bbox, i])

    finalContours = sorted(finalContours, key=lambda x: x[3][1])

    if draw:
        for con in finalContours:
            x, y, w, h = con[3]
            cv2.rectangle(imgDraw, (x, y), (x+w, y+h), (255, 0, 255), 3)
    return imgDraw, finalContours



def getRoi(contours):
    """
    Extracts the regions of interest (ROIs) from the given contours.

    This function takes the contours detected by `getContours` and extracts the bounding boxes
    for each contour, returning a list of regions of interest (ROIs) as tuples.

    Args:
        contours (list): A list of contours, where each contour is represented as a list 
                          containing contour details like number of vertices, area, and bounding box.

    Returns:
        list: A list of tuples, where each tuple represents the coordinates of a bounding box (x1, y1, x2, y2).
    """
    roiList = []
    for con in contours:
        x, y, w, h = con[3]
        roiList.append((x, y, x+w, y+h))
    return roiList


def get_text_from_bbox(doc, page_number, bbox):
    """
    Extracts text from a specific region of a page within a PDF document.

    This function extracts the text from the area defined by a bounding box (`bbox`) on a specific page 
    of the PDF document.

    Args:
        doc (fitz.Document): The PDF document object.
        page_number (int): The page number (1-based).
        bbox (tuple): The bounding box coordinates (x1, y1, x2, y2) from which text will be extracted.

    Returns:
        str: The extracted text from the specified bounding box.
    """
    # Load the specified page (note: page_number is 1-based)
    page = doc.load_page(page_number - 1)
    
    # Get text from the specified bounding box area
    text = page.get_text("text", clip=fitz.Rect(bbox))
    return text


def roiDisplay(roiList, show_process=False):
    """
    Displays the regions of interest (ROIs) as cropped images.

    This function takes a list of regions of interest (ROIs) and optionally displays them as separate 
    cropped images for visual inspection.

    Args:
        roiList (list): A list of tuples representing bounding box coordinates (x1, y1, x2, y2).
        show_process (bool, optional): If True, the cropped images will be displayed. Default is False.

    Returns:
        None
    """
    for x, roi in enumerate(roiList):
        roi = cv2.resize(roi, (0, 0), None, 2, 2)
        cv2.imshow(f"Cropped Image {x}", roi) if show_process else ""


def save_to_txt(highlightedText, result_name, clear_after_wards=True, images_folder_path='.ospdf-tmp-images'):
    """
    Saves a list of highlighted text to a specified text file and optionally clears the images folder.

    Arguments:
    highlightedText -- A list of strings representing the highlighted text to be saved.
    result_name -- The name of the output text file. Defaults to "Result.txt" if not provided.
    clear_after_wards -- A boolean flag indicating whether to clear the images folder after saving. Defaults to True.
    images_folder_path -- Path to the images folder to clear, if applicable. Defaults to '.ospdf-tmp-images'.

    Functionality:
    - Writes each entry of `highlightedText` to a new line in the specified file.
    - Clears the images folder if `clear_after_wards` is True.

    Returns:
    None
    """
    with open(result_name or "Result.txt", 'w') as f:
        for text in highlightedText:
            f.writelines(f'\n{text}')
    if clear_after_wards:
        clear_images_folder(images_folder_path)


def stackImages(scale, imgArray):
    """
    Stacks multiple images into a single display grid.

    Arguments:
    scale -- A scaling factor to resize the images (e.g., 0.5 for half size).
    imgArray -- A nested list or array of images. Each inner list represents a row.

    Functionality:
    - Resizes images to match the dimensions of the first image in the grid.
    - Converts grayscale images to RGB.
    - Stacks the images horizontally and vertically into a single composite.

    Returns:
    A single image representing the stacked grid.
    """
    rows, cols = len(imgArray), len(imgArray[0])

    rowsAvailable = isinstance(imgArray[0], list)
    width = imgArray[0][0].shape[1]
    height = imgArray[0][0].shape[0]
    if rowsAvailable:
        for x in range(0, rows):
            for y in range(0, cols):
                if imgArray[x][y].shape[:2] == imgArray[0][0].shape[:2]:
                    imgArray[x][y] = cv2.resize(imgArray[x][y], (0, 0), None, scale, scale)
                else: imgArray[x][y] = cv2.resize(imgArray[x][y], (imgArray[0][0].shape[1], imgArray[0][0].shape[0]), None, scale, scale)

                if len(imgArray[x][y].shape) == 2:
                    imgArray[x][y] = cv2.cvtColor(imgArray[x][y], cv2.COLOR_BGR2RGB)
        imageBlank = np.zeros((height, width, 3), np.uint8)
        hor = [imageBlank] * rows
        hor_copy = [imageBlank] * rows
        for x in range(0, rows):
            hor[x] = np.hstack(imgArray[x])
        ver = np.vstack(hor)
    else:
        for x in range(0, rows):
            if imgArray[x].shape[:2] == imgArray[0].shape[:2]:
                imgArray[x] = cv2.resize(imgArray[x], (0, 0), None, scale, scale)
            else:
                imgArray[x] = cv2.resize(imgArray[x], (imgArray[0].shape[1], imgArray[0].shape[0]), None, scale, scale)
            if len(imgArray[x].shape) == 2:
                imgArray[x] = cv2.cvtColor(imgArray[x], cv2.COLOR_BGR2BGRA)
        hor = np.hstack(imgArray)
        ver = hor
    return ver


def split_pdf(input_pdf, output_path, start_page=None, end_page=None):
    """
    Splits a PDF into a single page or a range of pages and saves the result.

    Arguments:
    input_pdf -- Path to the input PDF file.
    output_path -- Path to save the resulting PDF.
    start_page -- The starting page number (1-indexed). Defaults to None.
    end_page -- The ending page number (1-indexed). Defaults to None.

    Functionality:
    - If `start_page` is provided but `end_page` is not, extracts a single page.
    - If both `start_page` and `end_page` are provided, extracts the range of pages.
    - Ensures the page numbers are valid and within the document's range.

    Returns:
    None
    """

    # Open the input PDF
    doc = fitz.open(input_pdf)

    # Handle single page extraction
    if start_page is not None and end_page is None:
        # Create a new PDF for the single page
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=start_page - 1, to_page=start_page - 1)
        
        # Save the new PDF with the page number
        new_doc.save(output_path)
        new_doc.close()
        doc.close()
        return 

    # If both start_page and end_page are provided, extract the range of pages
    if start_page is not None and end_page is not None:
        # Ensure page numbers are within the valid range
        start_page = max(start_page - 1, 0)  # fitz uses zero-indexed pages
        end_page = min(end_page - 1, doc.page_count - 1)

        # Create a new PDF for the range of pages
        new_doc = fitz.open()
        
        # Insert pages in the specified range
        new_doc.insert_pdf(doc, from_page=start_page, to_page=end_page)

        # Save the new PDF with the page range
        new_doc.save(output_path)
        new_doc.close()
        doc.close()
        return 


def pdf_page_to_image(doc, page_number):    
    """
    Converts a specific page of a PDF document to an OpenCV-compatible image.

    Arguments:
    doc -- The `fitz.Document` object representing the PDF.
    page_number -- The page number (1-indexed) to convert to an image.

    Functionality:
    - Loads the specified page and converts it to a pixmap.
    - Converts the pixmap to a NumPy array in BGR format for OpenCV.

    Returns:
    A NumPy array representing the page as an image, or None if the page number is invalid.
    """
    # Ensure the page number is valid
    if page_number < 1 or page_number > doc.page_count:
        print("Invalid page number!")
        return None
    
    # Get the specified page (0-indexed, so subtract 1)
    page = doc.load_page(page_number - 1)
    
    # Convert the page to a pixmap (image)
    pix = page.get_pixmap()
    
    # Convert the pixmap to a PIL image
    pil_image = Image.open(io.BytesIO(pix.tobytes()))
    rgb_array = np.array(pil_image)

    # Convert RGB to BGR (OpenCV uses BGR format)
    bgr_image = cv2.cvtColor(rgb_array, cv2.COLOR_RGB2BGR)

    return bgr_image


def save_image(image, output_image_path):
    """
    Saves an image to a specified file path.

    Arguments:
    image -- The image to save, either as a NumPy array (OpenCV format) or a PIL Image object.
    output_image_path -- The file path where the image should be saved.

    Functionality:
    - Converts the image to PIL format if it's a NumPy array.
    - Saves the image to the specified location.

    Returns:
    None
    """
    if image is not None:
        # Convert numpy array (OpenCV format) to PIL.Image if needed
        if isinstance(image, np.ndarray):
            image = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))  # Convert BGR to RGB
        image.save(output_image_path)
        print(f"Image saved to {output_image_path}")
    else:
        print("No image to save.")



def save_to_pdf(results, filename, images_folder_path='.ospdf-tmp-images', clear_after_wards=True):
    """
    Saves text and images as a formatted PDF document.

    Arguments:
    results -- A list containing text, headings, and image file paths to include in the PDF.
    filename -- The name of the output PDF file.
    images_folder_path -- Path to the folder containing images, if applicable. Defaults to '.ospdf-tmp-images'.
    clear_after_wards -- A boolean flag to clear the images folder after saving. Defaults to True.

    Functionality:
    - Formats text and images into a styled PDF.
    - Handles headings, bullet points, and image captions.
    - Clears the images folder if `clear_after_wards` is True.

    Returns:
    None
    """
    # Create the PDF document
    doc = SimpleDocTemplate(filename, pagesize=letter)
    elements = []

    # Get default styles
    styles = getSampleStyleSheet()

    # Custom Heading style (using a large font and bold)
    heading_style = ParagraphStyle(
        name="Heading1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        textColor=colors.darkblue,
        spaceAfter=12,
        alignment=0,  # Left Centered ?
    )

    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=12,
        textColor=colors.black,
        spaceAfter=10,
        bulletFontName="Helvetica",
        bulletFontSize=50,
        bulletColor=colors.black,
        leftIndent=25,
        bulletSymbol="◉",
    )

    # Iterate over the results and add them as paragraphs
    i = 0
    while i < len(results):
        item = results[i]

        # Check if the item is a heading or normal text
        if "#######" in item:
            heading = item.split("#######")[1].split("##############")[0].strip()
            elements.append(Paragraph(heading, heading_style))
        else:
            # Handle Image Embedding
            if os.path.isfile(item):
                # Embed the image
                img_path = item
                img_width = 200
                img_height = 200
                img = Img(img_path, width=img_width, height=img_height)  # Correct way to instantiate Image class
                elements.append(img)

                # Check if the next line is a caption for the image
                if i + 1 < len(results) and "Figure" in results[i + 1]:
                    caption = results[i + 1]
                    caption_style = getSampleStyleSheet()["Normal"]
                    elements.append(Paragraph(f"<i>{caption}</i>", caption_style))
                    i += 1  # Skip the next item since it's the caption

            elif item == "\n\n==============================================================\n\n":
                elements.append(Spacer(1, 12))  # Adds space before the line

                # Add a horizontal line
                line = HRFlowable(width="100%", color=colors.black, thickness=1)
                elements.append(line)
                elements.append(Spacer(1, 12))  # Add space after the line
            elif re.match(r"^Page \d+:$", item):
                create_header_footer(elements, item)
            elif re.match(r"\n\nOriginal Text Length \d+\n' {6}\| {6}\nHighlighted Text Count: \d+", item):
                create_header_footer(elements, item, is_footer=True)
            else:
                # Add bullet points or other text
                elements.append(Paragraph(f"◉   {item}", bullet_style))

        i += 1

    # Build the PDF
    doc.build(elements)

    # Optional: Clear the images folder if specified
    if clear_after_wards:
        clear_images_folder(images_folder_path)


def detectImages(img):
    """
    Detects and highlights regions of interest (e.g., objects) in an image.

    Arguments:
    img -- The input image in BGR format.

    Functionality:
    - Converts the image to grayscale.
    - Detects edges and contours.
    - Filters contours by area and aspect ratio to identify regions of interest.
    - Draws rectangles around detected regions and returns cropped regions.

    Returns:
    A tuple containing:
    - The original image with rectangles drawn around detected regions.
    - A list of cropped image regions.
    """
    # Convert to grayscale
    imgGray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Apply edge detection
    edges = cv2.Canny(imgGray, 50, 150)

    # Dilate edges to connect nearby components
    kernel = np.ones((5, 5), np.uint8)
    imgDilated = cv2.dilate(edges, kernel, iterations=1)

    # Find contours
    contours, _ = cv2.findContours(imgDilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    imgContours = img.copy()
    imageRegions = []

    for contour in contours:
        area = cv2.contourArea(contour)
        if area > 1000:  # Adjust the area threshold based on your use case
            x, y, w, h = cv2.boundingRect(contour)
            aspectRatio = w / float(h)

            # Filter based on aspect ratio or size if needed
            if 0.5 < aspectRatio < 3:  # Adjust aspect ratio range
                cv2.rectangle(imgContours, (x, y), (x + w, y + h), (255, 0, 255), 2)
                imageRegions.append(img[y:y+h, x:x+w])  # Crop image regions

    return imgContours, imageRegions



def contains_number_pattern(string, caption=False):
    """
    Checks if a given string matches a specific number pattern.

    Arguments:
    string -- The string to be checked for patterns.
    caption -- A boolean flag to determine whether to match a caption pattern.
               If True, matches figure references like "Figure 1.1". Defaults to False.

    Functionality:
    - Normalizes the input string by replacing newline characters with spaces.
    - If `caption` is True, checks for patterns indicating figure references.
    - If `caption` is False, checks for general heading patterns while excluding figure references.

    Returns:
    A boolean indicating whether the string matches the specified pattern.
    """
    caption_pattern = r".*\bFigure \d+\.\d+(\.\d+)?\b.*"  # Matches figure references like "Figure 1.1"
    general_pattern = r"^\s*\d+\.\d+(\.\d+)?\b.*"  # Matches valid headings (e.g., "1.1 Introduction")

    string = string.replace("\n", " ")  # Normalize string for pattern matching

    if caption:
        # Check for caption pattern
        return bool(re.match(caption_pattern, string))

    # Exclude figure references
    if not re.search(caption_pattern, string):
        # Check for general heading pattern
        return bool(re.match(general_pattern, string))

    return False



def getTextFromPDFAsParagraphs(doc, page_number):
    """
    Extracts text from a specific page of a PDF document and splits it into paragraphs.

    Arguments:
    doc -- The `fitz.Document` object representing the PDF.
    page_number -- The page number (1-indexed) from which to extract text.

    Functionality:
    - Loads the specified page and extracts its text.
    - Splits the text into two forms of paragraphs:
      - `paragraphsForNormal`: Normalized by replacing newlines with spaces.
      - `paragraphsForHeadings`: Retains newline characters for accurate heading extraction.

    Returns:
    A tuple containing:
    - `paragraphsForNormal`: A list of cleaned paragraphs.
    - `paragraphsForHeadings`: A list of raw paragraphs for heading analysis.
    """
    text = doc.load_page(page_number-1).get_text()
    paragraphsForNormal = [i.replace("\n", " ").strip() for i in text.split('. ') if i.replace("\n", " ").strip()]
    paragraphsForHeadings = [i.strip() for i in text.split('. ') if i.strip()]
    return paragraphsForNormal, paragraphsForHeadings



def displayResult(title, iterable):
    """
    Displays the content of an iterable with a given title.

    Arguments:
    title -- The title to display above the content.
    iterable -- An iterable containing the content to display.

    Functionality:
    - Prints the title followed by the contents of the iterable.
    - Each item in the iterable is printed on a new line.

    Returns:
    None
    """
    print(f"\n\n\n\n\n\n\n\n{title}: \n\n")
    for i in iterable: 
        print(i, end="\n\n")



def get_count(iterable, counter):
    """
    Counts the total number of items in a nested iterable.

    Arguments:
    iterable -- A nested iterable (e.g., list of lists) to count items in.
    counter -- An integer to store the running total count.

    Functionality:
    - Iterates through the outer iterable and counts the items in each inner element.

    Returns:
    The total count of items as an integer.
    """

    for i in iterable:
        counter += len(i)
    return counter



def create_header_footer(elements, item, is_footer=False, page_width=letter[0], page_height=letter[1]):
    """
    Adds a header or footer to a PDF document.

    Arguments:
    elements -- A list to which the header or footer elements are added.
    item -- The text content for the header or footer.
    is_footer -- A boolean indicating whether the element is a footer (True) or a header (False).
                 Defaults to False.
    page_width -- The width of the page in points. Defaults to `letter[0]`.
    page_height -- The height of the page in points. Defaults to `letter[1]`.

    Functionality:
    - Styles the header or footer with predefined `ParagraphStyle`.
    - Adds appropriate spacing before or after the element.
    - Appends the header or footer content to the `elements` list.

    Returns:
    None
    """
    # If it's a header or a footer
    header_style = ParagraphStyle(
        name="Header",
        fontName="Courier",
        fontSize=10,
        textColor=colors.grey,
        spaceAfter=6,
        alignment=0
    )

    footer_style = ParagraphStyle(
        name="Footer",
        fontName="Courier",
        fontSize=10,
        textColor=colors.grey,
        spaceBefore=6,
        alignment=2
    )
    
    if is_footer:
        # Add footer with page number or other footer text
        elements.append(Spacer(1, 12))  # Spacer before footer
        elements.append(Paragraph(item, footer_style))  
        elements.append(Spacer(1, 12))  

    else:
        # Add header
        elements.append(Spacer(1, 12))  
        elements.append(Paragraph(item, header_style))  
        elements.append(Spacer(1, 12))  





def load_images_from_folder(folder):
    """
    Loads images from a specified folder and returns them as a list.

    Arguments:
    + folder -- Path to the folder containing image files.

    Functionality:
    - Iterates through files in the folder.
    - Filters files by image extensions (e.g., .png, .jpg, .jpeg).
    - Reads valid image files into OpenCV format and stores their paths.

    Returns:
    A tuple containing:
    - `images`: A list of loaded images in OpenCV format.
    - `paths`: A list of corresponding file paths for the loaded images.
    """

    images = []
    paths = []
    for filename in os.listdir(folder):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
            img_path = os.path.join(folder, filename)
            img = cv2.imread(img_path)
            if img is not None:
                images.append(img)
                paths.append(img_path)
    return images, paths





def mouse_callback(event, x, y, flags, param):
    """
    Handles mouse events for interacting with images in a grid.

    Arguments:
    + event -- The type of mouse event (e.g., `cv2.EVENT_LBUTTONDOWN`).
    + x -- The x-coordinate of the mouse click.
    + y -- The y-coordinate of the mouse click.
    + flags -- Additional flags related to the mouse event.
    + param -- A dictionary containing:
             - `grid_coords`: A list of tuples defining the grid's rows, columns, and image paths.
             - `img_paths`: A list of paths to images in the grid.
             - `close_flag`: A boolean flag to signal closing the interaction.
             - `pdf_image_path`: A string to store the selected image path.

    Functionality:
    - Checks if the click is within the bounds of any grid cell.
    - If so, updates `pdf_image_path` with the clicked image's path.
    - Sets `close_flag` to True to indicate the interaction is complete.

    Returns:
    None
    """
    if event == cv2.EVENT_LBUTTONDOWN:
        grid_coords = param["grid_coords"]
        img_paths = param["img_paths"]
        close_flag = param["close_flag"]
        pdf_image_path = param["pdf_image_path"]
        for (row, col, path) in grid_coords:
            if row[0] <= y <= row[1] and col[0] <= x <= col[1]:
                print(f"Clicked on image: {path}")
                pdf_image_path = path  # Update the pdf_image_path
                close_flag = True  # Set close flag to True

        param["close_flag"] = close_flag
        param["pdf_image_path"] = pdf_image_path





def display_images_grid(images, paths, window_name="Image Grid", cell_size=(200, 200), padding=10, close_image_window=False, pdf_image_path=''):
    """
    Displays multiple images in a grid format within an OpenCV window.

    Arguments:
    + images -- A list of image objects (numpy arrays) to display.
    + paths -- A list of file paths corresponding to the images.
    + window_name -- Name of the window (default: "Image Grid").
    + cell_size -- A tuple (width, height) specifying the size of each grid cell (default: (200, 200)).
    + padding -- Spacing between grid cells (default: 10 pixels).
    + close_image_window -- Boolean flag to indicate whether to close the window programmatically (default: False).
    + pdf_image_path -- A string to store the path of the selected image (default: '').

    Functionality:
    - Arranges the given images into a grid layout with specified cell size and padding.
    - Displays the grid in a resizable OpenCV window.
    - Allows mouse interaction for selecting images within the grid.
    - Dynamically determines grid dimensions based on the number of images.

    Returns:
    + A string containing the path of the selected image or an empty string if no selection is made.
    """

    if not images:
        print("No images to display.")
        return ''
    
    print(f"Total images to display: {len(images)}")
    
    rows = int(np.ceil(len(images) ** 0.5))  # Determine grid size
    cols = rows

    grid_height = max(10, rows * cell_size[1] + (rows - 1) * padding)
    grid_width = max(10, cols * cell_size[0] + (cols - 1) * padding)

    background_color = (75, 0, 130)  # RGB for light blue
    grid_img = np.full((grid_height, grid_width, 3), background_color, dtype=np.uint8)


    # grid_img = np.ones((grid_height, grid_width, 3), dtype=np.uint8) * 255  # White background for grid

    grid_coords = []  # List to store the coordinates for each image cell
    for idx, img in enumerate(images):
        if img is None:
            print(f"Error: Image at {paths[idx]} could not be loaded.")
            continue
        resized_img = cv2.resize(img, cell_size)  # Resize the image to fit in the cell
        row_idx = idx // cols
        col_idx = idx % cols

        y_start = row_idx * (cell_size[1] + padding)
        x_start = col_idx * (cell_size[0] + padding)
        y_end = y_start + cell_size[1]
        x_end = x_start + cell_size[0]

        grid_img[y_start:y_end, x_start:x_end] = resized_img  # Place image in the grid
        grid_coords.append(((y_start, y_end), (x_start, x_end), paths[idx]))  # Store image coordinates and path

    if not grid_coords:
        print("No images were placed in the grid.")
        return ''
    
    # Set up the window and mouse callback
    cv2.namedWindow(window_name)
    param = {"grid_coords": grid_coords, "img_paths": paths, "close_flag": close_image_window, "pdf_image_path": pdf_image_path}
    cv2.setMouseCallback(window_name, mouse_callback, param)

    while True:
        cv2.imshow(window_name, grid_img)
        key = cv2.waitKey(1)
        if key == 27 or param["close_flag"]:  # Press ESC to close or if flag is True
            break

    cv2.destroyAllWindows()
    return param["pdf_image_path"]


def string_mouse_callback(event, x, y, flags, param):
    """
    Handles mouse events for selecting strings in an OpenCV window.

    Arguments:
    + event -- The type of mouse event (e.g., `cv2.EVENT_LBUTTONDOWN`).
    + x -- The x-coordinate of the mouse click.
    + y -- The y-coordinate of the mouse click.
    + flags -- Additional flags related to the mouse event.
    + param -- A dictionary containing:
            - `strings`: A list of strings displayed in the window.
            - `close_flag`: Boolean flag to close the window.
            - `pdf_caption_text`: The selected string for further use.

    Functionality:
    - Detects if a click occurs on a string in the displayed list.
    - Updates `pdf_caption_text` with the selected string.
    - Sets `close_flag` to True to signal that the selection is complete.

    Returns:
    None
    """

    if event == cv2.EVENT_LBUTTONDOWN:
        strings = param["strings"]
        close_flag = param["close_flag"]
        pdf_caption_text = param["pdf_caption_text"]
        y_offset = 50  # Vertical offset for the text lines
        text_height = 30  # Height of each text line
        click_padding = 10  # Additional padding between text lines for easier clicking
        clicked_text = None
        
        # Check if the click is within the bounds of any string
        for idx, text in enumerate(strings):
            y_start = y_offset + idx * (text_height + click_padding)  # Add padding to each line's position
            y_end = y_start + text_height
            # If the click position (y) is within the range of a string's position
            if y_start <= y <= y_end:
                clicked_text = text
                break
        
        if clicked_text:
            print(f"Clicked on: {clicked_text}")
            pdf_caption_text = clicked_text  # Update the caption text
            close_flag = True  # Set close flag to True

        param["close_flag"] = close_flag
        param["pdf_caption_text"] = pdf_caption_text



def display_strings(strings, window_name="Strings List", close_caption_window = False, pdf_caption_text = ''):
    """
    Displays a list of strings in an OpenCV window for user interaction.

    Arguments:
    + strings -- A list of strings to display in the window.
    + window_name -- Name of the window (default: "Strings List").
    + close_caption_window -- Boolean flag to control programmatic closure of the window (default: False).
    + pdf_caption_text -- A string to store the selected caption (default: '').

    Functionality:
    - Renders the list of strings on a white background.
    - Allows mouse-based selection of a string from the list.
    - Returns the selected string for further use.

    Returns:
    + A string containing the selected caption text or an empty string if no selection is made.
    """

    # Create a blank image to draw the strings on
    img = np.zeros((500, 700, 3), dtype=np.uint8)  # Adjust size if needed
    img[:] = (255, 255, 255)  # White background

    y_offset = 50  # Start Y-position for the first string
    text_height = 30  # Height of each line of text
    click_padding = 10  # Additional padding between text lines for easier clicking

    # Draw each string onto the image
    for idx, text in enumerate(strings):
        y_position = y_offset + idx * (text_height + click_padding)
        cv2.putText(img, text, (50, y_position + text_height // 2), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 0, 0), 2, cv2.LINE_AA)

    # Set up the window and mouse callback
    cv2.namedWindow(window_name)
    param = {"strings": strings, "close_flag": close_caption_window, "pdf_caption_text": pdf_caption_text}
    cv2.setMouseCallback(window_name, string_mouse_callback, param)


    while True:
        cv2.imshow(window_name, img)
        key = cv2.waitKey(1)
        if key == 27 or param["close_flag"]:  # Press ESC to close or click on an option
            break

    cv2.destroyAllWindows()
    return param["pdf_caption_text"]


def clear_images_folder(folder_path):
    """
    Deletes all contents of a specified folder, including subdirectories and files.

    Arguments:
    + folder_path -- The path to the folder to be cleared.

    Functionality:
    - Checks if the folder exists.
    - Deletes all files and subdirectories within the folder.
    - Prints a confirmation message upon successful deletion or a warning if the folder does not exist.

    Returns:
    None
    """

    # Check if folder exists
    if os.path.exists(folder_path):
        # Loop through all files and subdirectories
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            
            # If it's a directory, remove it recursively
            if os.path.isdir(file_path):
                shutil.rmtree(file_path)
            else:
                # If it's a file, remove it
                os.remove(file_path)
        print(f"All contents of {folder_path} have been deleted.")
    else:
        print(f"The folder {folder_path} does not exist.")


def merge_pdfs(output_pdf_file, *input_pdf_files):
    """
    Merges multiple PDF files into a single output PDF.

    Arguments:
    + output_pdf_file -- The path to the resulting merged PDF file.
    + *input_pdf_files -- A variable number of input PDF file paths to be merged.

    Functionality:
    - Creates a new PDF document for the merged output.
    - Iterates over the input PDF files, appending their pages to the merged PDF.
    - Handles exceptions, printing an error message if a file cannot be opened or merged.
    - Saves the final merged PDF to the specified output path.

    Returns:
    None
    """

    # Create a new PDF document to store the merged result
    merged_pdf = fitz.open()

    # Loop through each input PDF and append its pages to the merged PDF
    for pdf_file in input_pdf_files:
        try:
            with fitz.open(pdf_file) as pdf:
                merged_pdf.insert_pdf(pdf)
        except Exception as e:
            print(f"Error: Unable to open or merge the file '{pdf_file}'. {e}")
            return

    # Save the merged PDF to the output file
    merged_pdf.save(output_pdf_file)
    print(f"Successfully merged PDFs into {output_pdf_file}")




def create_header_footer_for_docx(doc, item, is_footer=False):
    """
    Adds a header or footer to a Word document with custom formatting.

    Arguments:
    + doc -- The Word document object to which the header/footer is added.
    + item -- The text to include in the header or footer.
    + is_footer -- Boolean flag indicating whether the text is for a footer (default: False).

    Functionality:
    - Defines a style for the header/footer, including font, size, color, and alignment.
    - Creates a paragraph in the document with the given text, applying the defined style.
    - Adjusts spacing based on whether it is a header or footer.

    Returns:
    None
    """

    header_footer_style = {
        "font_name": "Courier",
        "font_size": 10,
        "color": RGBColor(128, 128, 128),  # Grey color
        "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT if not is_footer else WD_PARAGRAPH_ALIGNMENT.RIGHT,
    }

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(item)
    run.font.name = header_footer_style["font_name"]
    run.font.size = Pt(header_footer_style["font_size"])
    run.font.color.rgb = header_footer_style["color"]
    paragraph.alignment = header_footer_style["alignment"]

    # Add spacing
    if is_footer:
        paragraph.paragraph_format.space_before = Pt(6)
    else:
        paragraph.paragraph_format.space_after = Pt(6)





def save_to_docx(results, filename, images_folder_path='.ospdf-tmp-images', clear_after_wards=True):
    """
    Saves a list of results into a Word document with formatting and optional images.

    Arguments:
    + results -- A list of strings or file paths to add to the document. Supports headings, text, images, and captions.
    + filename -- The path to save the resulting Word document.
    + images_folder_path -- The path to the folder containing images to embed (default: '.ospdf-tmp-images').
    + clear_after_wards -- Boolean flag to delete the contents of `images_folder_path` after saving the document (default: True).

    Functionality:
    - Creates a Word document with custom styles for headings, bullet points, and embedded images.
    - Iterates over the `results` list, identifying and formatting different types of content:
    - Headings: Recognized by delimiters `#######` and `##############`.
    - Images: Embedded with specified dimensions and optional captions.
    - Horizontal lines: Rendered for specific markers.
    - Headers and footers: Added for specific patterns.
    - Bullet points: Applied to regular text items.
    - Optionally clears the images folder after saving the document.

    Returns:
    None
    """

    # Create the Word document
    doc = Document()

    # Custom Heading style
    heading_style = {
        "font_size": 15,
        "bold": True,
        "color": RGBColor(0, 0, 139),  # Dark blue
        "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
    }

    # Custom Bullet style
    bullet_style = {
        "font_size": 12,
        "color": RGBColor(0, 0, 0),  # Black
        "indentation": 25,
        "symbol": "◉",
    }

    # Iterate over the results and add them to the document
    i = 0
    while i < len(results):
        item = results[i]

        # Check if the item is a heading or normal text
        if "#######" in item:
            heading = item.split("#######")[1].split("##############")[0].strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(heading)
            run.font.size = Pt(heading_style["font_size"])
            run.bold = heading_style["bold"]
            run.font.color.rgb = heading_style["color"]
            paragraph.alignment = heading_style["alignment"]
        else:
            # Handle Image Embedding
            if os.path.isfile(item):
                # Embed the image
                img_path = item
                doc.add_picture(img_path, width=Pt(200), height=Pt(200))

                # Check if the next line is a caption for the image
                if i + 1 < len(results) and "Figure" in results[i + 1]:
                    caption = results[i + 1]
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(caption)
                    run.font.size = Pt(12)
                    run.italic = True
                    i += 1  # Skip the next item since it's the caption

            elif item == "\n\n==============================================================\n\n":
                # Add a horizontal line
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("=" * 50)
                run.font.size = Pt(12)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif re.match(r"^Page \d+:$", item):
                create_header_footer_for_docx(doc, item)
            elif re.match(r"\n\nOriginal Text Length \d+\n' {6}\| {6}\nHighlighted Text Count: \d+", item):
                create_header_footer_for_docx(doc, item, is_footer=True)
            else:
                # Add bullet points or other text
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{bullet_style['symbol']}   {item}")
                run.font.size = Pt(bullet_style["font_size"])
                run.font.color.rgb = bullet_style["color"]
                paragraph.paragraph_format.left_indent = Pt(bullet_style["indentation"])

        i += 1

    # Save the Word document
    doc.save(filename)

    # Optional: Clear the images folder if specified
    if clear_after_wards:
        clear_images_folder(images_folder_path)



def gradient_text(text, colors):
    """
    Apply a gradient color effect to the input text.

    This function takes a string of text and a list of three colors, and returns
    the text with a gradient effect applied from left to right. The gradient will
    transition smoothly between the provided colors, based on the position of each
    character in the text.

    Args:
    - text (str): The input text to which the gradient will be applied.
    - colors (list of int): A list of three integers representing color codes
      for the gradient. The first color represents the start of the gradient,
      the second is the middle color, and the third is the end color.

    Returns:
    - str: The input text with the gradient applied to each character.
    
    Example:
    >>> gradient_text("Hello World", [196, 214, 226])
    'Text with gradient colors applied'
    """
    start_color = colors[0]  
    middle_color = colors[1]  
    end_color = colors[2]  
    text_lines = text.splitlines()
    
    def interpolate_color(start, end, factor):
        """
        Interpolate between two color values based on a given factor.
        
        Args:
        - start (int): The start color value.
        - end (int): The end color value.
        - factor (float): A factor between 0 and 1 indicating the interpolation point.

        Returns:
        - int: The interpolated color value.
        """
        return int(start + (end - start) * factor)
    
    colored_text = []
    for line in text_lines:
        colored_line = ""
        for i, char in enumerate(line):
            factor = i / len(line)
            if factor < 0.5:
                color_code = interpolate_color(start_color, middle_color, factor * 2)
            else:
                color_code = interpolate_color(middle_color, end_color, (factor - 0.5) * 2)
            colored_char = f"\033[38;5;{color_code}m{char}\033[0m"
            colored_line += colored_char
        colored_text.append(colored_line)
    return "\n".join(colored_text)



def check_and_create_folder(folder_name):
    # Get the current working directory
    current_directory = os.getcwd()
    
    # Define the full path to the folder
    folder_path = os.path.join(current_directory, folder_name)
    
    # Check if the folder exists
    if not os.path.exists(folder_path):
        # Create the folder if it doesn't exist
        os.makedirs(folder_path)
        print(f"Folder '{folder_name}' created.")