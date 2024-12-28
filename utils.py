import cv2
import os
import numpy as np
import re
import fitz
from PIL import Image
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.platypus import Image as Img
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.pdfgen import canvas
import time
import shutil
import threading
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json










STATE_FILE = "state.json"

def save_state(pdf_path, persist_state):
    state = {
        "pdf_path": pdf_path,
        "persist_state": persist_state
    }
    with open(STATE_FILE, "w") as f:
        json.dump(state, f)  

def load_state():
    try:
        with open(STATE_FILE, "r") as f:
            state = json.load(f)
        pdf_path = state.get("pdf_path")
        persist_state = state.get("persist_state")
        return pdf_path, persist_state
    except FileNotFoundError:
        raise ValueError("State file not found.")
    except json.JSONDecodeError:
        raise ValueError("Error decoding state file.")

def clear_state():
    if os.path.exists(STATE_FILE):
        os.remove(STATE_FILE)
        print("\nCurrent State Cleared\n")






def detectColor(img, hsv, show_process=False):
    imgHSV = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
    cv2.imshow("HSV", imgHSV) if show_process else ""
    lower = np.array([hsv[0], hsv[2], hsv[4]])
    upper = np.array([hsv[1], hsv[3], hsv[5]])
    mask = cv2.inRange(imgHSV, lower, upper)
    imgResult = cv2.bitwise_and(img, img, mask=mask)
    cv2.imshow("Result", imgResult) if show_process else ""
    return imgResult, imgHSV


def getContours(img, imgDraw, showCanny=False, minArea=1000, filter=0, cThr=[100, 100], draw=True):
    imgDraw = imgDraw.copy()
    imgGray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    imgBlur = cv2.GaussianBlur(imgGray, (5, 5), 1)
    imgCanny = cv2.Canny(imgBlur, cThr[0], cThr[1])
    kernel = np.ones((10, 10), np.uint8)
    imgDial = cv2.dilate(imgCanny, kernel, iterations=1)
    imgClose = cv2.morphologyEx(imgDial, cv2.MORPH_CLOSE, kernel)

    if showCanny: cv2.imshow("Canny", imgClose)

    contours, hierarchy = cv2.findContours(imgClose, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    finalContours = []
    for i in contours:
        area = cv2.contourArea(i)
        if area > minArea:
            peri = cv2.arcLength(i, True)
            approx = cv2.approxPolyDP(i, 0.02*peri, True)
            bbox = cv2.boundingRect(approx)
            if filter > 0: 
                if len(approx) == filter:
                    finalContours.append([len(approx), area, approx, bbox, i])
            else:
                finalContours.append([len(approx), area, approx, bbox, i])

    finalContours = sorted(finalContours, key=lambda x: x[3][1])

    if draw:
        for con in finalContours:
            x, y, w, h = con[3]
            cv2.rectangle(imgDraw, (x, y), (x+w, y+h), (255, 0, 255), 3)
    return imgDraw, finalContours



def getRoi(contours):
    roiList = []
    for con in contours:
        x, y, w, h = con[3]
        roiList.append((x, y, x+w, y+h))
    return roiList


def get_text_from_bbox(doc, page_number, bbox):
    # Load the specified page (note: page_number is 1-based)
    page = doc.load_page(page_number - 1)
    
    # Get text from the specified bounding box area
    text = page.get_text("text", clip=fitz.Rect(bbox))
    return text


def roiDisplay(roiList, show_process=False):
    for x, roi in enumerate(roiList):
        roi = cv2.resize(roi, (0, 0), None, 2, 2)
        cv2.imshow(f"Cropped Image {x}", roi) if show_process else ""


def saveText(highlightedText, result_name):
    with open(result_name or "Result.txt", 'w') as f:
        for text in highlightedText:
            f.writelines(f'\n{text}')



def stackImages(scale, imgArray):
    rows, cols = len(imgArray), len(imgArray[0])

    rowsAvailable = isinstance(imgArray[0], list)
    width = imgArray[0][0].shape[1]
    height = imgArray[0][0].shape[0]
    if rowsAvailable:
        for x in range(0, rows):
            for y in range(0, cols):
                if imgArray[x][y].shape[:2] == imgArray[0][0].shape[:2]:
                    imgArray[x][y] = cv2.resize(imgArray[x][y], (0, 0), None, scale, scale)
                else: imgArray[x][y] = cv2.resize(imgArray[x][y], (imgArray[0][0].shape[1], imgArray[0][0].shape[0]), None, scale, scale)

                if len(imgArray[x][y].shape) == 2:
                    imgArray[x][y] = cv2.cvtColor(imgArray[x][y], cv2.COLOR_BGR2RGB)
        imageBlank = np.zeros((height, width, 3), np.uint8)
        hor = [imageBlank] * rows
        hor_copy = [imageBlank] * rows
        for x in range(0, rows):
            hor[x] = np.hstack(imgArray[x])
        ver = np.vstack(hor)
    else:
        for x in range(0, rows):
            if imgArray[x].shape[:2] == imgArray[0].shape[:2]:
                imgArray[x] = cv2.resize(imgArray[x], (0, 0), None, scale, scale)
            else:
                imgArray[x] = cv2.resize(imgArray[x], (imgArray[0].shape[1], imgArray[0].shape[0]), None, scale, scale)
            if len(imgArray[x].shape) == 2:
                imgArray[x] = cv2.cvtColor(imgArray[x], cv2.COLOR_BGR2BGRA)
        hor = np.hstack(imgArray)
        ver = hor
    return ver


def split_pdf(input_pdf, output_folder, start_page=None, end_page=None):
    # Open the input PDF
    doc = fitz.open(input_pdf)

    # Handle single page extraction
    if start_page is not None and end_page is None:
        # Create a new PDF for the single page
        new_doc = fitz.open()
        new_doc.insert_pdf(doc, from_page=start_page - 1, to_page=start_page - 1)
        
        # Save the new PDF with the page number
        new_doc.save(f"{output_folder}/page_{start_page}.pdf")
        new_doc.close()
        doc.close()
        return 

    # If both start_page and end_page are provided, extract the range of pages
    if start_page is not None and end_page is not None:
        # Ensure page numbers are within the valid range
        start_page = max(start_page - 1, 0)  # fitz uses zero-indexed pages
        end_page = min(end_page - 1, doc.page_count - 1)

        # Create a new PDF for the range of pages
        new_doc = fitz.open()
        
        # Insert pages in the specified range
        new_doc.insert_pdf(doc, from_page=start_page, to_page=end_page)

        # Save the new PDF with the page range
        new_doc.save(f"{output_folder}/pages_{start_page + 1}_to_{end_page + 1}.pdf")
        new_doc.close()
        doc.close()
        return 


def pdf_page_to_image(doc, page_number):    
    # Ensure the page number is valid
    if page_number < 1 or page_number > doc.page_count:
        print("Invalid page number!")
        return None
    
    # Get the specified page (0-indexed, so subtract 1)
    page = doc.load_page(page_number - 1)
    
    # Convert the page to a pixmap (image)
    pix = page.get_pixmap()
    
    # Convert the pixmap to a PIL image
    pil_image = Image.open(io.BytesIO(pix.tobytes()))
    rgb_array = np.array(pil_image)

    # Convert RGB to BGR (OpenCV uses BGR format)
    bgr_image = cv2.cvtColor(rgb_array, cv2.COLOR_RGB2BGR)

    return bgr_image


def save_image(image, output_image_path):
    """
    Save a PIL Image object to a file.
    
    :param image: The PIL Image object to save.
    :param output_image_path: The file path where the image should be saved.
    """
    if image:
        image.save(output_image_path)
        print(f"Image saved to {output_image_path}")
    else:
        print("No image to save.")



def save_to_pdf(results, filename, images_folder_path='images/images', clear_after_wards=True):
    # Create the PDF document
    doc = SimpleDocTemplate(filename, pagesize=letter)
    elements = []

    # Get default styles
    styles = getSampleStyleSheet()

    # Custom Heading style (using a large font and bold)
    heading_style = ParagraphStyle(
        name="Heading1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        textColor=colors.darkblue,
        spaceAfter=12,
        alignment=0,  # Left Centered ?
    )

    bullet_style = ParagraphStyle(
        name="Bullet",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=12,
        textColor=colors.black,
        spaceAfter=10,
        bulletFontName="Helvetica",
        bulletFontSize=50,
        bulletColor=colors.black,
        leftIndent=25,
        bulletSymbol="◉",
    )

    # Iterate over the results and add them as paragraphs
    i = 0
    while i < len(results):
        item = results[i]

        # Check if the item is a heading or normal text
        if "#######" in item:
            heading = item.split("#######")[1].split("##############")[0].strip()
            elements.append(Paragraph(heading, heading_style))
        else:
            # Handle Image Embedding
            if os.path.isfile(item):
                # Embed the image
                img_path = item
                img_width = 200
                img_height = 200
                img = Img(img_path, width=img_width, height=img_height)  # Correct way to instantiate Image class
                elements.append(img)

                # Check if the next line is a caption for the image
                if i + 1 < len(results) and "Figure" in results[i + 1]:
                    caption = results[i + 1]
                    caption_style = getSampleStyleSheet()["Normal"]
                    elements.append(Paragraph(f"<i>{caption}</i>", caption_style))
                    i += 1  # Skip the next item since it's the caption

            elif item == "\n\n==============================================================\n\n":
                elements.append(Spacer(1, 12))  # Adds space before the line

                # Add a horizontal line
                line = HRFlowable(width="100%", color=colors.black, thickness=1)
                elements.append(line)
                elements.append(Spacer(1, 12))  # Add space after the line
            elif re.match(r"^Page \d+:$", item):
                create_header_footer(elements, item)
            elif re.match(r"\n\nOriginal Text Length \d+\n' {6}\| {6}\nHighlighted Text Count: \d+", item):
                create_header_footer(elements, item, is_footer=True)
            else:
                # Add bullet points or other text
                elements.append(Paragraph(f"◉   {item}", bullet_style))

        i += 1

    # Build the PDF
    doc.build(elements)

    # Optional: Clear the images folder if specified
    if clear_after_wards:
        clear_images_folder(images_folder_path)


def detectImages(img):
    # Convert to grayscale
    imgGray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

    # Apply edge detection
    edges = cv2.Canny(imgGray, 50, 150)

    # Dilate edges to connect nearby components
    kernel = np.ones((5, 5), np.uint8)
    imgDilated = cv2.dilate(edges, kernel, iterations=1)

    # Find contours
    contours, _ = cv2.findContours(imgDilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    imgContours = img.copy()
    imageRegions = []

    for contour in contours:
        area = cv2.contourArea(contour)
        if area > 1000:  # Adjust the area threshold based on your use case
            x, y, w, h = cv2.boundingRect(contour)
            aspectRatio = w / float(h)

            # Filter based on aspect ratio or size if needed
            if 0.5 < aspectRatio < 3:  # Adjust aspect ratio range
                cv2.rectangle(imgContours, (x, y), (x + w, y + h), (255, 0, 255), 2)
                imageRegions.append(img[y:y+h, x:x+w])  # Crop image regions

    return imgContours, imageRegions



def contains_number_pattern(string, caption=False):
    caption_pattern = r".*\bFigure \d+\.\d+(\.\d+)?\b.*"  # Matches figure references like "Figure 1.1"
    general_pattern = r"^\s*\d+\.\d+(\.\d+)?\b.*"  # Matches valid headings (e.g., "1.1 Introduction")

    string = string.replace("\n", " ")  # Normalize string for pattern matching

    if caption:
        # Check for caption pattern
        return bool(re.match(caption_pattern, string))

    # Exclude figure references
    if not re.search(caption_pattern, string):
        # Check for general heading pattern
        return bool(re.match(general_pattern, string))

    return False



def getTextFromPDFAsParagraphs(doc, page_number):
    text = doc.load_page(page_number-1).get_text()
    paragraphsForNormal = [i.replace("\n", " ").strip() for i in text.split('. ') if i.replace("\n", " ").strip()]
    paragraphsForHeadings = [i.strip() for i in text.split('. ') if i.strip()]
    return paragraphsForNormal, paragraphsForHeadings



def displayResult(title, iterable):
    print(f"\n\n\n\n\n\n\n\n{title}: \n\n")
    for i in iterable: 
        print(i, end="\n\n")



def get_count(iterable, counter):
    for i in iterable:
        counter += len(i)
    return counter



def create_header_footer(elements, item, is_footer=False, page_width=letter[0], page_height=letter[1]):
    # If it's a header or a footer
    header_style = ParagraphStyle(
        name="Header",
        fontName="Courier",
        fontSize=10,
        textColor=colors.grey,
        spaceAfter=6,
        alignment=0
    )

    footer_style = ParagraphStyle(
        name="Footer",
        fontName="Courier",
        fontSize=10,
        textColor=colors.grey,
        spaceBefore=6,
        alignment=2
    )
    
    if is_footer:
        # Add footer with page number or other footer text
        elements.append(Spacer(1, 12))  # Spacer before footer
        elements.append(Paragraph(item, footer_style))  
        elements.append(Spacer(1, 12))  

    else:
        # Add header
        elements.append(Spacer(1, 12))  
        elements.append(Paragraph(item, header_style))  
        elements.append(Spacer(1, 12))  




# Function to load images from a folder
def load_images_from_folder(folder):
    images = []
    paths = []
    for filename in os.listdir(folder):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
            img_path = os.path.join(folder, filename)
            img = cv2.imread(img_path)
            if img is not None:
                images.append(img)
                paths.append(img_path)
    return images, paths




# Mouse callback function to handle clicks
def mouse_callback(event, x, y, flags, param):
    if event == cv2.EVENT_LBUTTONDOWN:
        grid_coords = param["grid_coords"]
        img_paths = param["img_paths"]
        close_flag = param["close_flag"]
        pdf_image_path = param["pdf_image_path"]
        for (row, col, path) in grid_coords:
            if row[0] <= y <= row[1] and col[0] <= x <= col[1]:
                print(f"Clicked on image: {path}")
                pdf_image_path = path  # Update the pdf_image_path
                close_flag = True  # Set close flag to True

        param["close_flag"] = close_flag
        param["pdf_image_path"] = pdf_image_path




# Function to display images in a grid
def display_images_grid(images, paths, window_name="Image Grid", cell_size=(200, 200), padding=10, close_image_window=False, pdf_image_path=''):
    rows = int(np.ceil(len(images) ** 0.5))  # Determine grid size
    cols = rows

    grid_height = rows * cell_size[1] + (rows - 1) * padding
    grid_width = cols * cell_size[0] + (cols - 1) * padding

    background_color = (75, 0, 130)  # RGB for light blue
    grid_img = np.full((grid_height, grid_width, 3), background_color, dtype=np.uint8)


    # grid_img = np.ones((grid_height, grid_width, 3), dtype=np.uint8) * 255  # White background for grid

    grid_coords = []  # List to store the coordinates for each image cell
    for idx, img in enumerate(images):
        resized_img = cv2.resize(img, cell_size)  # Resize the image to fit in the cell
        row_idx = idx // cols
        col_idx = idx % cols

        y_start = row_idx * (cell_size[1] + padding)
        x_start = col_idx * (cell_size[0] + padding)
        y_end = y_start + cell_size[1]
        x_end = x_start + cell_size[0]

        grid_img[y_start:y_end, x_start:x_end] = resized_img  # Place image in the grid
        grid_coords.append(((y_start, y_end), (x_start, x_end), paths[idx]))  # Store image coordinates and path

    # Set up the window and mouse callback
    cv2.namedWindow(window_name)
    param = {"grid_coords": grid_coords, "img_paths": paths, "close_flag": close_image_window, "pdf_image_path": pdf_image_path}
    cv2.setMouseCallback(window_name, mouse_callback, param)

    while True:
        cv2.imshow(window_name, grid_img)
        key = cv2.waitKey(1)
        if key == 27 or param["close_flag"]:  # Press ESC to close or if flag is True
            break

    cv2.destroyAllWindows()
    return param["pdf_image_path"]


def string_mouse_callback(event, x, y, flags, param):
    if event == cv2.EVENT_LBUTTONDOWN:
        strings = param["strings"]
        close_flag = param["close_flag"]
        pdf_caption_text = param["pdf_caption_text"]
        y_offset = 50  # Vertical offset for the text lines
        text_height = 30  # Height of each text line
        click_padding = 10  # Additional padding between text lines for easier clicking
        clicked_text = None
        
        # Check if the click is within the bounds of any string
        for idx, text in enumerate(strings):
            y_start = y_offset + idx * (text_height + click_padding)  # Add padding to each line's position
            y_end = y_start + text_height
            # If the click position (y) is within the range of a string's position
            if y_start <= y <= y_end:
                clicked_text = text
                break
        
        if clicked_text:
            print(f"Clicked on: {clicked_text}")
            pdf_caption_text = clicked_text  # Update the caption text
            close_flag = True  # Set close flag to True

        param["close_flag"] = close_flag
        param["pdf_caption_text"] = pdf_caption_text



def display_strings(strings, window_name="Strings List", close_caption_window = False, pdf_caption_text = ''):

    # Create a blank image to draw the strings on
    img = np.zeros((500, 700, 3), dtype=np.uint8)  # Adjust size if needed
    img[:] = (255, 255, 255)  # White background

    y_offset = 50  # Start Y-position for the first string
    text_height = 30  # Height of each line of text
    click_padding = 10  # Additional padding between text lines for easier clicking

    # Draw each string onto the image
    for idx, text in enumerate(strings):
        y_position = y_offset + idx * (text_height + click_padding)
        cv2.putText(img, text, (50, y_position + text_height // 2), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 0, 0), 2, cv2.LINE_AA)

    # Set up the window and mouse callback
    cv2.namedWindow(window_name)
    param = {"strings": strings, "close_flag": close_caption_window, "pdf_caption_text": pdf_caption_text}
    cv2.setMouseCallback(window_name, string_mouse_callback, param)


    while True:
        cv2.imshow(window_name, img)
        key = cv2.waitKey(1)
        if key == 27 or param["close_flag"]:  # Press ESC to close or click on an option
            break

    cv2.destroyAllWindows()
    return param["pdf_caption_text"]


def clear_images_folder(folder_path):
    # Check if folder exists
    if os.path.exists(folder_path):
        # Loop through all files and subdirectories
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            
            # If it's a directory, remove it recursively
            if os.path.isdir(file_path):
                shutil.rmtree(file_path)
            else:
                # If it's a file, remove it
                os.remove(file_path)
        print(f"All contents of {folder_path} have been deleted.")
    else:
        print(f"The folder {folder_path} does not exist.")


def merge_pdfs(output_pdf_file, *input_pdf_files):
    # Create a new PDF document to store the merged result
    merged_pdf = fitz.open()

    # Loop through each input PDF and append its pages to the merged PDF
    for pdf_file in input_pdf_files:
        try:
            with fitz.open(pdf_file) as pdf:
                merged_pdf.insert_pdf(pdf)
        except Exception as e:
            print(f"Error: Unable to open or merge the file '{pdf_file}'. {e}")
            return

    # Save the merged PDF to the output file
    merged_pdf.save(output_pdf_file)
    print(f"Successfully merged PDFs into {output_pdf_file}")




def create_header_footer_for_docx(doc, item, is_footer=False):
    header_footer_style = {
        "font_name": "Courier",
        "font_size": 10,
        "color": RGBColor(128, 128, 128),  # Grey color
        "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT if not is_footer else WD_PARAGRAPH_ALIGNMENT.RIGHT,
    }

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(item)
    run.font.name = header_footer_style["font_name"]
    run.font.size = Pt(header_footer_style["font_size"])
    run.font.color.rgb = header_footer_style["color"]
    paragraph.alignment = header_footer_style["alignment"]

    # Add spacing
    if is_footer:
        paragraph.paragraph_format.space_before = Pt(6)
    else:
        paragraph.paragraph_format.space_after = Pt(6)





def save_to_docx(results, filename, images_folder_path='images/images', clear_after_wards=True):
    # Create the Word document
    doc = Document()

    # Custom Heading style
    heading_style = {
        "font_size": 15,
        "bold": True,
        "color": RGBColor(0, 0, 139),  # Dark blue
        "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
    }

    # Custom Bullet style
    bullet_style = {
        "font_size": 12,
        "color": RGBColor(0, 0, 0),  # Black
        "indentation": 25,
        "symbol": "◉",
    }

    # Iterate over the results and add them to the document
    i = 0
    while i < len(results):
        item = results[i]

        # Check if the item is a heading or normal text
        if "#######" in item:
            heading = item.split("#######")[1].split("##############")[0].strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(heading)
            run.font.size = Pt(heading_style["font_size"])
            run.bold = heading_style["bold"]
            run.font.color.rgb = heading_style["color"]
            paragraph.alignment = heading_style["alignment"]
        else:
            # Handle Image Embedding
            if os.path.isfile(item):
                # Embed the image
                img_path = item
                doc.add_picture(img_path, width=Pt(200), height=Pt(200))

                # Check if the next line is a caption for the image
                if i + 1 < len(results) and "Figure" in results[i + 1]:
                    caption = results[i + 1]
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(caption)
                    run.font.size = Pt(12)
                    run.italic = True
                    i += 1  # Skip the next item since it's the caption

            elif item == "\n\n==============================================================\n\n":
                # Add a horizontal line
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("=" * 50)
                run.font.size = Pt(12)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif re.match(r"^Page \d+:$", item):
                create_header_footer_for_docx(doc, item)
            elif re.match(r"\n\nOriginal Text Length \d+\n' {6}\| {6}\nHighlighted Text Count: \d+", item):
                create_header_footer_for_docx(doc, item, is_footer=True)
            else:
                # Add bullet points or other text
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"{bullet_style['symbol']}   {item}")
                run.font.size = Pt(bullet_style["font_size"])
                run.font.color.rgb = bullet_style["color"]
                paragraph.paragraph_format.left_indent = Pt(bullet_style["indentation"])

        i += 1

    # Save the Word document
    doc.save(filename)

    # Optional: Clear the images folder if specified
    if clear_after_wards:
        clear_images_folder(images_folder_path)

